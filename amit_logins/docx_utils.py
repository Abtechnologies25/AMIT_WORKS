import io
import os
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def generate_tax_invoice_docx(invoice):
    doc = Document()
    
    # Set page margins
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.77)
    section.right_margin = Cm(0.77)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    logo_path = os.path.join(settings.BASE_DIR, "static", "images", "logo.png")

    # We use a single table with 7 columns for perfect border alignment
    table = doc.add_table(rows=0, cols=7)
    table.style = 'Table Grid'
    table.autofit = False

    # Set table-level cell margins (in twips: 1 inch = 1440 twips)
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    for side, val in [('top', 40), ('start', 80), ('bottom', 40), ('end', 80)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)
    
    # Column widths (Total ~7.27 inches for A4 with 0.5" margins)
    # Col 0: S.NO         (0.6)  } PREPARED BY = 0+1 (2.0)
    # Col 1: DESC 1       (1.4)  } DESC = 1+2+3 (3.9) 
    # Col 2: DESC 2/SEAL  (1.9)  } SEAL = 2 (1.9) } SELLER = 0+1+2 (3.9)
    # Col 3: DESC 3/BILL1 (0.6)  } BILL = 3+4+5+6 (3.37)
    # Col 4: QTY          (0.5)  } 
    # Col 5: UNIT PRICE   (1.1)  } FOR AMIT = 3+4+5+6 (3.37)
    # Col 6: TOTAL VALUE  (1.17) }
    widths = [Inches(0.5), Inches(1.1), Inches(2.1), Inches(1.1), Inches(0.5), Inches(1.1), Inches(0.87)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    def _set_widths(row):
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    def add_p(cell, text, bold=False, size=10, align=None, space_after=0):
        if not cell.paragraphs[0].text and len(cell.paragraphs) == 1:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3 + space_after)
        p.paragraph_format.line_spacing = 1.0
        if align is not None:
            p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
        return p

    def _set_cell_vmerge(cell, val='restart'):
        """Apply vertical merge to a cell. val='restart' for first row, 'continue' for subsequent."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for existing in tcPr.findall(qn('w:vMerge')):
            tcPr.remove(existing)
        vMerge = OxmlElement('w:vMerge')
        if val == 'restart':
            vMerge.set(qn('w:val'), 'restart')
        tcPr.append(vMerge)

    def _set_cell_margins(cell, top=40, start=80, bottom=40, end=80):
        """Override cell-level margins (in twips). 1 inch = 1440 twips."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    row1 = table.add_row()
    _set_widths(row1)
    # Logo cell: merge cols 0+1 horizontally
    c_logo_r1 = row1.cells[0]
    c_logo_r1.merge(row1.cells[1])
    c_logo_r1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Logo cell: tight left/right, small top/bottom (matches image)
    _set_cell_margins(c_logo_r1, top=30, start=5, bottom=30, end=5)
    if os.path.exists(logo_path):
        p_logo = c_logo_r1.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(0)
        p_logo.add_run().add_picture(logo_path, width=Cm(2.88), height=Cm(2.22))
    # TAX INVOICE: merge cols 2-6
    c_title = row1.cells[2]
    for i in range(3, 7):
        c_title.merge(row1.cells[i])
    add_p(c_title, "TAX INVOICE", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    c_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Row 2: Logo continued (col 0-1, vMerge) | INVOICE NO (col 2) | DATE (col 3-6) ---
    row2_h = table.add_row()
    _set_widths(row2_h)
    date_str = invoice.DATE.strftime('%d-%m-%Y') if invoice.DATE else ''
    # Logo continuation: merge cols 0+1 horizontally then set vMerge continue
    c_logo_r2 = row2_h.cells[0]
    c_logo_r2.merge(row2_h.cells[1])
    # INVOICE NO: col 2 only (aligns with SELLER boundary)
    c_inv = row2_h.cells[2]
    p_inv = add_p(c_inv, "INVOICE NO : ", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p_inv.add_run(str(invoice.INVOICE_NO)).font.size = Pt(10)
    c_inv.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # DATE: cols 3-6 (aligns with BILL boundary)
    c_date = row2_h.cells[3]
    c_date.merge(row2_h.cells[4])
    c_date.merge(row2_h.cells[5])
    c_date.merge(row2_h.cells[6])
    p_date = add_p(c_date, "DATE : ", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p_date.add_run(date_str).font.size = Pt(10)
    c_date.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Apply vertical merge so logo spans both rows
    _set_cell_vmerge(c_logo_r1, 'restart')
    _set_cell_vmerge(c_logo_r2, 'continue')

    # 3a. Seller and Bill Titles
    row = table.add_row()
    _set_widths(row)
    c0 = row.cells[0]
    c0.merge(row.cells[1])
    c0.merge(row.cells[2])
    
    c3 = row.cells[3]
    c3.merge(row.cells[4])
    c3.merge(row.cells[5])
    c3.merge(row.cells[6])
    
    add_p(c0, "SELLER", bold=True, size=11)
    add_p(c3, "BUYER", bold=True, size=11)
    # BILL title cell: add left padding
    _set_cell_margins(c3, top=40, start=80, bottom=40, end=80)
    
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 3b. Seller and Bill Addresses
    row2 = table.add_row()
    _set_widths(row2)
    c0_addr = row2.cells[0]
    c0_addr.merge(row2.cells[1])
    c0_addr.merge(row2.cells[2])
    
    c3_addr = row2.cells[3]
    c3_addr.merge(row2.cells[4])
    c3_addr.merge(row2.cells[5])
    c3_addr.merge(row2.cells[6])
    # BILL address cell: add left padding matching SELLER
    _set_cell_margins(c3_addr, top=40, start=80, bottom=40, end=80)
    
    # Seller
    add_p(c0_addr, "ATHITH MITHRA INDUSTRIAL TECHNOLOGIES", bold=True, size=10)
    add_p(c0_addr, "(AMIT) PVT LTD", bold=True, size=10)
    add_p(c0_addr, "No.7/2, SLB School South Road,", size=10)
    add_p(c0_addr, "Ramavarmapuram, Nagercoil,", size=10)
    add_p(c0_addr, "Kanyakumari District,", size=10)
    add_p(c0_addr, "Tamil Nadu, India - 629 001.", size=10)
    
    # Aligned Contact, GST, CIN
    for lbl, val in [("Contact", "+91 9840511458"), ("GST", "33AAYCA9919P1ZK"), ("CIN", "U26517TN2023PTC161676")]:
        p = add_p(c0_addr, "", size=10)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.8))
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.95))
        
        run_l = p.add_run(lbl)
        run_l.bold = True
        run_l.font.size = Pt(10)
        
        run_c = p.add_run("\t:\t")
        run_c.bold = True
        run_c.font.size = Pt(10)
        
        run_v = p.add_run(val)
        run_v.font.size = Pt(10)

    # Bill
    bill_to_text = str(invoice.BILL_TO or "").strip()
    if bill_to_text:
        bill_lines = bill_to_text.split('\n')
        for idx, line in enumerate(bill_lines):
            line = line.strip()
            if line:
                # Make the first line bold (Company Name) if it looks like one, 
                # or just follow the same pattern as before.
                is_bold = (idx == 0)
                add_p(c3_addr, line, bold=is_bold, size=10)
    
    for cell in row2.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 4. Items Header
    row = table.add_row()
    _set_widths(row)
    c1 = row.cells[1]
    c1.merge(row.cells[2])
    c1.merge(row.cells[3])
    
    headers = ["S.NO", "DESCRIPTION", "QTY", "UNIT PRICE", "TOTAL VALUE"]
    cells = [row.cells[0], row.cells[1], row.cells[4], row.cells[5], row.cells[6]]
    for idx, text in enumerate(headers):
        add_p(cells[idx], text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Items Data
    items = invoice.items.all()
    for item in items:
        row = table.add_row()
        row.allow_break_across_pages = False
        _set_widths(row)
        c1 = row.cells[1]
        c1.merge(row.cells[2])
        c1.merge(row.cells[3])
        cells = [row.cells[0], row.cells[1], row.cells[4], row.cells[5], row.cells[6]]
        
        add_p(cells[0], str(item.S_NO), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(cells[1], str(item.DESCRIPTION or ""))
        add_p(cells[2], str(item.QTY or ""), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(cells[3], f"₹ {item.UNIT_PRICE}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_p(cells[4], f"₹ {item.TOTAL_VALUE}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    # 5. Summary Rows
    def add_summary_row(label, value, bold=False, size=10):
        row = table.add_row()
        row.allow_break_across_pages = False
        _set_widths(row)
        c0 = row.cells[0]
        for i in range(1, 6):
            c0.merge(row.cells[i])
        add_p(c0, label, bold=bold, size=size, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_p(row.cells[6], value, bold=bold, size=size, align=WD_ALIGN_PARAGRAPH.RIGHT)
        
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    add_summary_row("TOTAL AMOUNT", f"₹ {invoice.TOTAL_AMOUNT}")
    add_summary_row("GST 18%", f"₹ {invoice.GST_18}")
    add_summary_row("TOTAL AMOUNT WITH GST 18%", f"₹ {invoice.TOTAL_AMOUNT_WITH_GST}", bold=True, size=10)
    
    val = float(invoice.ROUND_OFF)
    prefix = "(-)" if val < 0 else "(+)"
    round_off_val = f"{prefix} {abs(val):.2f}"
    
    add_summary_row("ROUND OFF", round_off_val)
    add_summary_row("GRAND TOTAL", f"₹ {invoice.GRAND_TOTAL}", bold=True, size=10)

    # 6. Amount in Words
    row = table.add_row()
    _set_widths(row)
    c0 = row.cells[0]
    for i in range(1, 7):
        c0.merge(row.cells[i])
    
    p = c0.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run1 = p.add_run("AMOUNT IN WORDS: ")
    run1.bold = True
    run1.font.size = Pt(10)
    
    # Check if words already contains ONLY to avoid double ONLY
    words = str(invoice.AMOUNT_IN_WORDS).upper().strip()
    if not words.endswith("ONLY"):
        words = f"{words} ONLY"
    
    run2 = p.add_run(words)
    run2.font.size = Pt(10)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 7a. Bank Account & Terms Titles
    row = table.add_row()
    row.allow_break_across_pages = False
    _set_widths(row)
    c0 = row.cells[0]
    c0.merge(row.cells[1])
    c0.merge(row.cells[2])
    
    c3 = row.cells[3]
    c3.merge(row.cells[4])
    c3.merge(row.cells[5])
    c3.merge(row.cells[6])
    
    add_p(c0, "BANK ACCOUNT DETAILS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).keep_with_next = True
    add_p(c3, "TERMS & CONDITIONS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).keep_with_next = True
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 7b. Bank Account & Terms Content
    row_content = table.add_row()
    _set_widths(row_content)
    c0_c = row_content.cells[0]
    c0_c.merge(row_content.cells[1])
    c0_c.merge(row_content.cells[2])
    
    c3_c = row_content.cells[3]
    c3_c.merge(row_content.cells[4])
    c3_c.merge(row_content.cells[5])
    c3_c.merge(row_content.cells[6])
    
    bank_items = [
        ("Account Name", "ATHITH MITHRA"),
        ("", "INDUSTRIAL TECHNOLOGIES"),
        ("", "(AMIT) PVT LTD."),
        ("Bank Name", "CANARA BANK"),
        ("Account Number", "120026789274"),
        ("IFSC", "CNRB0001112"),
        ("Account Type", "CA"),
        ("Branch", "MEENAKSHIPURAM,"),
        ("", "NAGERCOIL.")
    ]
    
    for idx, (key, val) in enumerate(bank_items):
        p = add_p(c0_c, "", size=10)
        p.paragraph_format.space_before = Pt(6) if idx == 0 else Pt(0)
        p.paragraph_format.space_after = Pt(6) if idx == len(bank_items) - 1 else Pt(0)
        p.paragraph_format.left_indent = Inches(0.15)
        # Tab stops for aligning colon and value
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.6))
        if key:
            p.add_run(f"{key}\t:\t{val}").font.size = Pt(10)
        else:
            p.add_run(f"\t\t{val}").font.size = Pt(10)
        p.keep_with_next = True
            
    terms = [
        "Goods once sold will not be taken back.",
        "Delivery Charges should be in the hands of Customer.",
        "Any damages during Transport, report immediately in front of the Driver/ Delivery Person.",
        "Warranty norms are as per the guidelines."
    ]
    for idx, term in enumerate(terms):
        p = add_p(c3_c, "", size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        # Add bold bullet with non-breaking space to prevent stretching gap
        run_bullet = p.add_run("•\u00A0")
        run_bullet.bold = True
        run_bullet.font.size = Pt(10)
        # Add normal text
        run_text = p.add_run(term)
        run_text.bold = False
        run_text.font.size = Pt(10)
        
        p.paragraph_format.space_before = Pt(6) if idx == 0 else Pt(0)
        p.paragraph_format.space_after = Pt(6) if idx == len(terms) - 1 else Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)  # Text starts at 0.3
        p.paragraph_format.first_line_indent = Inches(-0.10) # Bullet starts at 0.3 - 0.15 = 0.15 (Matches Account Name)
        p.keep_with_next = True
        
    for cell in row_content.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 8. Signature Block (Part of the main table for perfect alignment)
    # The signature section rows use keep_with_next to stay together if possible

    # Row 8a: Titles
    row8a = table.add_row()
    row8a.allow_break_across_pages = False
    _set_widths(row8a)
    c0 = row8a.cells[0]
    c0.merge(row8a.cells[1])
    c2 = row8a.cells[2]
    c3 = row8a.cells[3]
    c3.merge(row8a.cells[4])
    c3.merge(row8a.cells[5])
    c3.merge(row8a.cells[6])
    
    add_p(c0, "PREPARED BY", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(c2, "SEAL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(c3, "FOR AMIT INDUSTRIAL\nTECHNOLOGIES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    for cell in row8a.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True

    # Row 8b: Blanks (Fixed height)
    row8b = table.add_row()
    row8b.allow_break_across_pages = False
    row8b.height = Inches(1.2) # Slightly larger height for the blank space
    _set_widths(row8b)
    c0 = row8b.cells[0]
    c0.merge(row8b.cells[1])
    c2 = row8b.cells[2]
    c3 = row8b.cells[3]
    c3.merge(row8b.cells[4])
    c3.merge(row8b.cells[5])
    c3.merge(row8b.cells[6])
    
    add_p(c0, "\n\n\n\n\n")
    add_p(c2, "\n\n\n\n\n")
    add_p(c3, "\n\n\n\n\n")
    
    for cell in row8b.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True

    # Row 8c: Footer (Merged full row)
    row8c = table.add_row()
    row8c.allow_break_across_pages = False
    _set_widths(row8c)
    # Merge the full row (Cols 0-6)
    c_footer = row8c.cells[0]
    for i in range(1, 7):
        c_footer.merge(row8c.cells[i])
    add_p(c_footer, "AUTHORISED SIGNATORY", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    c_footer.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in c_footer.paragraphs:
        p.paragraph_format.keep_together = True
    
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()
