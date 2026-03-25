import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        start={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        end={"sz": 12, "val": "single", "color": "000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_invoice_docx():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # 1. Header Table
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    cell_a1 = table.cell(0, 0)
    cell_b1 = table.cell(0, 2)
    cell_a1.merge(cell_b1)
    
    p = cell_a1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TAX INVOICE")
    run.bold = True
    run.font.size = Pt(20)
    
    cell_a2 = table.cell(1, 0)
    cell_b2 = table.cell(1, 1)
    cell_a2.merge(cell_b2)
    
    cell_c2 = table.cell(1, 2)
    p2 = cell_c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run("INVOICE NO : 044\nDATE : 16-03-2026")
    run2.bold = True
    run2.font.size = Pt(10)

    doc.save('test_invoice.docx')
    print("Test doc created")

if __name__ == "__main__":
    create_invoice_docx()
